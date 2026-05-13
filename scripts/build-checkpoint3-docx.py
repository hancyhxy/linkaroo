#!/usr/bin/env python3
"""
Render checkpoint3-submission.md content into the UTS-supplied
Checkpoint3Template_Updated.docx layout, preserving the template's
font scheme (Aptos / Aptos Display) and heading hierarchy.

Run:
    /tmp/docx-venv/bin/python aussiebridge/scripts/build-checkpoint3-docx.py

Output:
    uts-coursework/ios_innovation_studio/checkpoint3/
        Checkpoint3_AussieBridge_XinyiHan.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/qiansui/Desktop/xinyihan")
REPO = ROOT / "aussiebridge"
SHOTS = REPO / "screenshots"
MOCKUPS = REPO / "docs" / "mockups"
OUT = ROOT / "uts-coursework/ios_innovation_studio/checkpoint3/Checkpoint3_AussieBridge_XinyiHan.docx"

BODY_FONT = "Aptos"
HEAD_FONT = "Aptos Display"
SIZE_BODY = Pt(12)
SIZE_CAPTION = Pt(10)
SIZE_H3 = Pt(13)
SIZE_H2 = Pt(16)
SIZE_H1 = Pt(24)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_LABEL = RGBColor(0x0E, 0x28, 0x41)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)


def set_run(run, text, *, font=BODY_FONT, size=SIZE_BODY, bold=False, italic=False, color=COLOR_BLACK):
    run.text = text
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, *, font=BODY_FONT, size=SIZE_BODY, bold=False, italic=False, color=COLOR_BLACK, space_after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    set_run(p.add_run(), text, font=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level):
    sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(), text, font=HEAD_FONT, size=sizes[level], bold=True, color=COLOR_BLACK)
    return p


def add_label_value(doc, label, value, *, value_italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(), f"{label}  ", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)
    set_run(p.add_run(), value, font=BODY_FONT, size=SIZE_BODY, italic=value_italic)
    return p


def add_image(doc, path, *, width_inches=6.5, caption=None):
    """Insert a centered image with an optional italic caption beneath."""
    if not Path(path).exists():
        add_para(doc, f"[missing image: {path}]", italic=True, color=COLOR_MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap.paragraph_format.space_after = Pt(10)
        set_run(cap.add_run(), caption, font=BODY_FONT, size=SIZE_CAPTION, italic=True, color=COLOR_MUTED)


def set_cell_borders(cell):
    """Add subtle 1px borders so the inventory table reads as a real grid."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "B0B0B0")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_screenshot_grid(doc, rows):
    """rows = list of (img_path_left, caption_left, img_path_right, caption_right) tuples.
    Renders a 2-column borderless table; each cell has a centered image + caption.
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = False
    col_w = Inches(3.1)
    for r_idx, (img_l, cap_l, img_r, cap_r) in enumerate(rows):
        for c_idx, (img, cap) in enumerate([(img_l, cap_l), (img_r, cap_r)]):
            cell = table.cell(r_idx, c_idx)
            cell.width = col_w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            if img and Path(img).exists():
                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run()
                run.add_picture(str(img), width=Inches(2.6))
                if cap:
                    cp = cell.add_paragraph()
                    cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cp.paragraph_format.space_after = Pt(8)
                    set_run(cp.add_run(), cap, font=BODY_FONT, size=SIZE_CAPTION, italic=True, color=COLOR_MUTED)


def add_inventory_table(doc, header, rows):
    """Render the SCREENSHOTS.md image inventory as a Word-native table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = True
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        set_run(p.add_run(), h, font=BODY_FONT, size=Pt(10), bold=True, color=COLOR_LABEL)
        set_cell_borders(hdr_cells[i])
        # subtle header fill
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EEF1F5")
        tc_pr.append(shd)
    # Body rows
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            set_run(p.add_run(), val, font=BODY_FONT, size=Pt(10))
            set_cell_borders(cells[c_idx])


def build():
    doc = Document()

    # Tighten page margins so the banner & flow images can breathe at 6.5" wide
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = SIZE_BODY

    # ─── Cover ──────────────────────────────────────────────────────────
    add_heading(doc, "Value Branch", level=1)
    add_para(doc, "Context-First", font=HEAD_FONT, size=Pt(20), bold=True, color=COLOR_LABEL, space_after=12)

    add_label_value(doc, "Student Name", "Xinyi Han")
    add_label_value(doc, "Student ID", "25751470")

    # ─── 1. Overview ────────────────────────────────────────────────────
    add_heading(doc, "1. Overview", level=1)

    # Banner image FIRST (per Xinyi's note: "上来需要先把我们那张 Banner 图放进去")
    add_image(doc, SHOTS / "00-banner.png", width_inches=6.5,
              caption="Linkaroo: a context-first app where identity, location, and urgency act as a filter, not a feature.")

    # Para 1 — The team's app concept (Describe per brief)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(), "The team's app concept. ", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)
    set_run(p.add_run(),
        "Linkaroo is a forum-style platform built by our team for newcomers "
        "arriving in Sydney. We organise their early questions into precise topic "
        "categories, so they can find the right thread fast. When a thread is not "
        "enough, the app matches them with a volunteer for a 1-on-1 conversation. "
        "Most newcomer apps stop at information; ours commits to escalating into a "
        "real human exchange."
    )

    # Para 2 — My individual value branch (Explain per brief: builds on + diverges)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(), "My individual value branch: Context-First. ", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)
    set_run(p.add_run(),
        "On top of the team concept, my branch adds one rule. Every personalisation "
        "decision the app makes must be visible on the screen where it is made. Five "
        "mechanisms realise this:"
    )

    mechanisms = [
        ("Identity Awareness", "Onboarding captures language, visa status, location, and arrival duration once. This becomes a persistent filter, not a profile field."),
        ("Contextual Home Hub", "A 2×5 service grid plus a \"Recommended for You\" module. Items carry visible tags such as STUDENT MATCH or TOP ANSWER."),
        ("Granular Metadata Tags", "Four orthogonal tag dimensions on every Q&A post: Authority (SENIOR MATCH), Reliability (SOURCE: TIKTOK, UNVERIFIED), Timeliness (OLD LAW, NEW), Urgency (NEEDS ANSWER)."),
        ("Reason-Based Volunteer Matching", "Match scores are always paired with human-readable reasons. Same university, same language, recent arrival."),
        ("Context-Aware Chat", "The originating Q&A post auto-attaches as a card. One tap shares profile tags with the volunteer."),
    ]
    for i, (name, desc) in enumerate(mechanisms, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(18)
        set_run(p.add_run(), f"{i}. ", bold=True)
        set_run(p.add_run(), f"{name}. ", bold=True)
        set_run(p.add_run(), desc)

    add_para(doc,
        "The team treats user identity as a backend signal that quietly sorts content. "
        "Context-First lifts that signal to the foreground. A Q&A post carries a "
        "STUDENT MATCH chip so the reader sees why it surfaced. A volunteer card spells "
        "out the match reasons in plain text, not a score. The chat opens with the "
        "originating Q&A pinned at the top, so context is shown rather than assumed."
    )

    # Para 3 — Why I chose this direction (Why meaningful per brief)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(), "Why I chose this direction. ", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)
    set_run(p.add_run(),
        "Checkpoint 1 user research returned one finding that re-shaped the brief for "
        "me. Participant 1 named language as a meta-barrier. It is the layer above "
        "translation that decides whether a non-native reader can parse who said what, "
        "when, and whether to trust it. Invisible personalisation breaks for that "
        "audience. A non-native reader being shown a \"smart\" English recommendation "
        "has no way to audit whether the recommendation is sound. I chose Context-First "
        "because it gives that user something they can read at every decision point: a "
        "reason, a tag, a pinned context. It changes the unit of trust from \"the "
        "system knows what is good for me\" to \"I can see the basis on which the "
        "system thinks so.\""
    )

    # ─── 2. Iterations ──────────────────────────────────────────────────
    add_heading(doc, "2. Iterations", level=1)

    iterations = [
        {
            "title": "2.1. First Iteration: Google Stitch",
            "tools": "Markdown (written design brief) · Google Stitch (AI design tool, brief-to-visual generation) · Claude Code (Anthropic CLI, brief authoring assist)",
            "made": (
                "Before opening any visual tool, I authored a written Context-First design "
                "brief (docs/product-context.md) that fixed the five mechanisms (identity "
                "awareness, contextual home, granular metadata tags, reason-based volunteer "
                "matching, context-aware chat) and the app background. That Markdown brief "
                "became the input to Google Stitch. Stitch did not generate the concept. It "
                "translated my brief into a 9-page visual hypothesis covering Personalization, "
                "Onboarding, Home, Q&A, Q&A Scroll, Volunteer, Community, Message, and Chat. "
                "Stitch also produced an auto-generated design specification (typography, "
                "colour, spacing scales) which I copied verbatim into docs/design.md as the "
                "starting visual grammar for downstream iterations.\n\n"
                "The intentional sequence matters: brief first, then visual. Stitch is fast at "
                "producing competent screens, but the screens only carry whatever structure "
                "the brief specifies. The five Context-First mechanisms had to be defined as "
                "written design constraints before Stitch could express them visually. That "
                "meant visible credibility tags on every Q&A row, \"Why she's a match for "
                "you\" reasons on volunteer cards, and originating-post auto-attach in chat. "
                "Each of these appears in the Stitch output because they were already locked "
                "in product-context.md."
            ),
            "shaped": (
                "Carried forward from Checkpoint 1: the team-level Context-First mechanism "
                "statement and the meta-barrier user-research quote, both of which became the "
                "test for every visual decision (line-height for non-native readers, tag "
                "credibility hierarchy, visible match-reasons before scores). Discarded: any "
                "direction that treated language as a feature toggle rather than a "
                "meta-barrier."
            ),
            "influenced": (
                "The Checkpoint 1 user-research interview transcript (the meta-barrier quote, "
                "captured verbatim). The methodological choice to author the brief in Markdown "
                "first, rather than going straight to a visual tool, came from a commitment "
                "to spec-as-contract that would carry through every later iteration."
            ),
        },
        {
            "title": "2.2. Second Iteration: HTML Mockup via Figma + Claude Code",
            "tools": "Google Stitch (handoff) · Figma · Figma MCP (Model Context Protocol) · Claude Code",
            "made": (
                "I exported the Stitch-generated 9 page designs into Figma as an editable "
                "design source. Figma replaced Stitch as the working surface for any visual "
                "adjustment because Stitch outputs are not editable past initial generation. "
                "I then connected Figma to Claude Code through the Figma MCP server, which "
                "exposes the Figma file structure (frames, components, layers, tokens) as a "
                "queryable context for the AI. Claude Code read each page from the Figma file "
                "via MCP and generated 9 reproducible static HTML mockups, written into "
                "docs/mockups/*.html (personalization, homepage, community, qa, qa_scroll, "
                "volunteer, message, chat, profile).\n\n"
                "Why this iteration matters: the HTML mockups become the visual ground truth "
                "for the SwiftUI implementation in later iterations. They are reproducible "
                "outside Stitch (open in any browser), version-controllable as text rather "
                "than as pixel files, and serve as the authoritative reference for tokens, "
                "copy strings, and layout when writing the final Swift pages. Without this "
                "step the team would be locked to Stitch's closed output as the only visual "
                "spec."
            ),
            "shaped": (
                "Carried forward from Iter 1: every Stitch screen as the visual baseline, "
                "preserved 1-to-1 in the HTML output. Discarded: Stitch as the working "
                "surface. Once exported to Figma, the working source moved with it."
            ),
            "influenced": (
                "The need to make the design contract reviewable and version-controllable. A "
                "Stitch screenshot cannot be diffed, queried, or rebuilt programmatically. An "
                "HTML mockup can. The Figma MCP integration was the connector that made this "
                "automation possible without manually re-drawing each page."
            ),
        },
        {
            "title": "2.3. Third Iteration: Spec & Component Definition (Xcode Canvas Validation)",
            "tools": "Claude Code · Markdown (4-block spec template) · SwiftUI · Xcode Canvas (#Preview)",
            "made": (
                "With the HTML mockups locked as visual ground truth, I shifted from visual "
                "design to written spec and code-level building blocks. Three written specs "
                "decompose the design problem along orthogonal axes:\n\n"
                "• docs/design.md: visual grammar (colour, typography, elevation, the "
                "\"no-line\" rule)\n"
                "• docs/struct.md: 22-entity data model organised into 5 Context-First layers "
                "(zero page-behaviour pollution)\n"
                "• docs/spec.md: per-page 4-block template (Overview, Parameters, Actions, "
                "Layout) covering §1 to §9\n\n"
                "In parallel I built a 30-component SwiftUI library in "
                "ABDesignSystem/Sources/Components/AB*.swift so that spec.md regions resolve "
                "to concrete Swift types rather than free-form decisions. Each component "
                "carries a #Preview block, validated in Xcode Canvas before the component is "
                "considered ready for page composition. Canvas Validation is the mechanism "
                "that closes the gap between spec language and runnable Swift: if a component "
                "cannot render in #Preview from its spec definition alone, the spec is "
                "incomplete and gets amended.\n\n"
                "I used Claude Code throughout, but the abstraction boundaries (design tokens "
                "vs components vs pages, single-responsibility files, the four orthogonal "
                "spec docs) were design intentions I locked first. Claude Code drafted the "
                "content under those constraints."
            ),
            "shaped": (
                "Carried forward from Iter 2: the HTML mockups as visual ground truth. Every "
                "design.md token and every component visual was traced back to a mockup file. "
                "Discarded: relying on Stitch's auto-generated copy text. Every string was "
                "re-authored against the user-research vocabulary (e.g., \"Match a "
                "volunteer\" rather than Stitch's generic \"Get help\"). Added: a Components/ "
                "layer and a struct.md data model, neither of which existed in Stitch or in "
                "the HTML mockups."
            ),
            "influenced": (
                "The realisation that the HTML mockups, while visually authoritative, did not "
                "bind a data model or a reusable component vocabulary. struct.md had to be "
                "authored separately and traced into the design.md vocabulary. This is when "
                "the four-doc separation became deliberate, and when Components stopped being "
                "an implementation detail and became a first-class design surface."
            ),
        },
        {
            "title": "2.4. Fourth Iteration: Xcode App Render & Integration",
            "tools": "Claude Code · Xcode 16 · XcodeGen (project.yml to .xcodeproj) · iOS 17 Simulator (iPhone 16 Pro) · SwiftUI standard libs (NavigationStack, @Observable, .environment)",
            "made": (
                "Iteration 4 turns the spec, mockups, and component library into a runnable "
                "iOS app. It happens in two consecutive sub-phases that together form a "
                "single iteration: page-by-page rendering, then app-level integration.\n\n"
                "Page-by-page render. I built each of the 9 pages in Sources/Pages/ one at a "
                "time, composing AB* components from Iter 3. Every page was matched to its "
                "HTML mockup from Iter 2 and to its 4-block spec from Iter 3. OnboardingView, "
                "HomeView, CommunityView, QAListView, QADetailView, VolunteerMatchView, "
                "MessageListView, ChatView, and ProfileView each ship as a self-contained "
                "SwiftUI view that renders in Xcode Canvas independently.\n\n"
                "App integration. Once the 9 pages each rendered standalone, I stood up the "
                "LinkarooApp/ host target via XcodeGen (declarative project.yml). The "
                "integration adds three concerns the page-level work could not supply: "
                "(a) AppState.swift, an @Observable single source of truth holding the "
                "current ABUser and each tab's navigation history; (b) Routing/Route.swift, a "
                "typed enum listing every from-here-to-there as a case, paired with "
                "navigationDestination(for: Route.self) for value-driven routing; (c) four "
                "Tabs/ wrappers (Home, Community, Messages, Profile), each owning its own "
                "NavigationStack(path:). Cross-tab handoff is wired end-to-end: the §5 to §6 "
                "to §8 escalation flow hops Home tab to Messages tab and back, and the §1 to "
                "§9 identity round-trip (Onboarding writes ABUser, Profile reads it back, "
                "Edit profile re-enters Onboarding prefilled) flows through this layer.\n\n"
                "Apple HIG alignment. The app uses native iOS patterns: SF Symbols for every "
                "glyph (chevron.up/down, bubble.left, square.and.arrow.up, plus.circle.fill, "
                "arrow.up, with 8+ usages across components like ABTag, ABButton, "
                "ABQAPostCard, ABTextInput). Value-driven NavigationStack routing per HIG iOS "
                "16+ guidance. Standard TabView for primary navigation. Tonal layering and "
                "no-line section breaks per HIG depth recommendations. Accessibility surface "
                "is partial: ABTabBar carries .accessibilityLabel and "
                ".accessibilityAddTraits(.isSelected) for tab state. docs/design.md calls "
                "for body-lg line-height of 1.6× explicitly to support non-native English "
                "readers. Full Dynamic Type and VoiceOver coverage are not yet wired across "
                "every component. This is a knowingly deferred axis."
            ),
            "shaped": (
                "Carried forward from Iter 3: every spec doc, every AB* component, every "
                "design token, untouched. Pages compose components rather than re-implement "
                "them. Discarded: nothing. The 4-layer architecture (tokens, components, "
                "models, pages) survived intact. App integration is purely additive: a host "
                "shell + shared state + routing layered above Pages."
            ),
            "influenced": (
                "The Checkpoint 3 rubric requirement to demonstrate \"essential core "
                "features that effectively demonstrate intended functionality for user "
                "stories.\" And the Context-First mechanism in product-context.md "
                "(shared-context handoff) only makes sense when state crosses pages, which "
                "only happens with an AppState and routing. So the App Target was not "
                "decoration. It is the only level at which the Context-First story is "
                "actually testable end-to-end."
            ),
        },
    ]

    for it in iterations:
        add_heading(doc, it["title"], level=2)
        add_label_value(doc, "Tools Used", it["tools"])

        for label, body in [
            ("What was Made", it["made"]),
            ("How the previous iteration shaped this", it["shaped"]),
            ("What/Who Influenced this Decision", it["influenced"]),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            set_run(p.add_run(), label, font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)
            for chunk in body.split("\n\n"):
                add_para(doc, chunk.strip(), space_after=6)

        # Iteration-level images intentionally omitted; user will place
        # screenshots manually after the docx is generated.

    # ─── 3. Final Prototype ─────────────────────────────────────────────
    add_heading(doc, "3. Final Prototype", level=1)
    add_heading(doc, "3.1. Final Prototype", level=2)

    add_label_value(doc, "GitHub Repository *", "https://github.com/hancyhxy/linkaroo")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(), "What was Made", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)

    add_para(doc,
        "A runnable iOS 17 SwiftUI app organised as two parallel pieces in one repository:"
    )

    final_pieces = [
        ("ABDesignSystem/", "a Swift Package providing 4 design-token files (ABColors, ABTypography, ABSpacing, ABElevation), 32 reusable AB* components, the 22-entity data model with mock data, and 9 graduated Pages. Every Component carries a #Preview block; the Package builds standalone via swift build."),
        ("LinkarooApp/", "a host iOS App Target generated declaratively via XcodeGen from project.yml. It depends on ABDesignSystem as a local SwiftPM package and adds the orchestration layer: @main entry, @Observable AppState, a Route enum, an onboarding gate, and 4 NavigationStack-backed tabs (Home / Community / Messages / Profile)."),
    ]
    for name, desc in final_pieces:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(18)
        set_run(p.add_run(), f"• {name}: ", bold=True)
        set_run(p.add_run(), desc)

    add_para(doc,
        "The project layers, top-down: design tokens, components, models, pages, "
        "routing/tabs, app entry. No third-party UI libraries; only Apple's SwiftUI and "
        "SwiftPM standard surface. AI tooling: I authored a written design brief "
        "(docs/product-context.md) first, then used Google Stitch to translate that "
        "brief into the 9-page visual hypothesis (Iteration 1). From Iteration 2 onward "
        "I used Claude Code (Anthropic CLI) as the primary AI tool for spec authoring "
        "(the 4-block template), prototype generation under the SPEC GAP discipline, "
        "gap-driven refactor (e.g. ABBackBar / ABPageHero split), access-control "
        "upgrade, routing wire-up, and the §1 to §9 single-source identity capture "
        "refactor. Every commit message names what changed and why. The git history is "
        "the audit trail for AI engagement."
    )

    # ── End-to-end flow diagram ──
    add_heading(doc, "End-to-End Navigation Flow", level=3)
    add_para(doc,
        "The full app navigation graph: 9 Pages spread across 4 tabs (Home, Community, "
        "Messages, Profile), with cross-tab handoff for the §5 to §6 to §8 escalation "
        "flow and the §1 to §9 identity round-trip."
    )
    add_image(doc, SHOTS / "flow_linkaroo.png", width_inches=6.5,
              caption="Complete navigation graph generated from the LinkarooApp Routing layer.")

    # ── 9-page screenshot catalogue ──
    add_heading(doc, "9-Page Screenshot Catalogue", level=3)
    add_para(doc,
        "Every spec-driven page in §1 → §9 order. Pages with a meaningful second state "
        "(scrolled view, alternate step) carry both states inline so the rhythm is visible."
    )
    grid_rows = [
        (SHOTS / "01-onboarding.png", "§1 Onboarding · top",
         SHOTS / "01b-onboarding-step2.png", "§1 Onboarding · step 2"),
        (SHOTS / "02-home.png", "§2 Home · top",
         SHOTS / "02b-home-scrolled.png", "§2 Home · Recommended for You"),
        (SHOTS / "03-community.png", "§3 Community · carousel",
         SHOTS / "03b-community-scrolled.png", "§3 Community · Q&A list"),
        (SHOTS / "04-qa-list.png", "§4 Q&A List",
         SHOTS / "04b-qa-list-scrolled.png", "§4 Q&A List · scrolled"),
        (SHOTS / "05-qa-detail.png", "§5 Q&A Detail · top",
         SHOTS / "05b-qa-detail-scrolled.png", "§5 Q&A Detail · answers"),
        (SHOTS / "06-volunteer-match.png", "§6 Volunteer Match · Top Choice",
         SHOTS / "06b-volunteer-match-scrolled.png", "§6 Volunteer Match · More matches"),
        (SHOTS / "07-message-list.png", "§7 Message List",
         SHOTS / "08-chat.png", "§8 Chat · shared context card"),
        (SHOTS / "09-profile.png", "§9 Profile · captured identity",
         None, None),
    ]
    add_screenshot_grid(doc, grid_rows)

    # ── Image inventory table (SCREENSHOTS.md bottom) ──
    add_heading(doc, "Image Inventory", level=3)
    add_para(doc,
        "Slot-locked filenames map each design surface to its captured screenshot, "
        "supporting the python-docx pipeline that generated this document."
    )
    inventory_header = ["Slot", "Filename", "Status"]
    inventory_rows = [
        ("Banner (Overview)", "screenshots/00-banner.png", "filled"),
        ("§1 Onboarding", "screenshots/01-onboarding.png", "filled"),
        ("§1b Onboarding step 2", "screenshots/01b-onboarding-step2.png", "filled"),
        ("§2 Home", "screenshots/02-home.png", "filled"),
        ("§2b Home scrolled", "screenshots/02b-home-scrolled.png", "filled"),
        ("§3 Community", "screenshots/03-community.png", "filled"),
        ("§3b Community scrolled", "screenshots/03b-community-scrolled.png", "filled"),
        ("§4 Q&A List", "screenshots/04-qa-list.png", "filled"),
        ("§4b Q&A List scrolled", "screenshots/04b-qa-list-scrolled.png", "filled"),
        ("§5 Q&A Detail", "screenshots/05-qa-detail.png", "filled"),
        ("§5b Q&A Detail scrolled", "screenshots/05b-qa-detail-scrolled.png", "filled"),
        ("§6 Volunteer Match", "screenshots/06-volunteer-match.png", "filled"),
        ("§6b Volunteer Match scrolled", "screenshots/06b-volunteer-match-scrolled.png", "filled"),
        ("§7 Message List", "screenshots/07-message-list.png", "filled"),
        ("§8 Chat", "screenshots/08-chat.png", "filled"),
        ("§9 Profile", "screenshots/09-profile.png", "filled"),
        ("Navigation flow diagram", "screenshots/flow_linkaroo.png", "filled"),
    ]
    add_inventory_table(doc, inventory_header, inventory_rows)

    # ── How the iterations led here ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(), "How the iterations led here", font=BODY_FONT, size=SIZE_BODY, bold=True, color=COLOR_LABEL)

    add_para(doc,
        "Iteration 1 produced the visual hypothesis and the Context-First brief in "
        "concrete pixels. Iteration 2 produced the structural contracts (design.md, "
        "struct.md, spec.md, Components) and separated what the system is from how each "
        "page expresses it. Iteration 3 pressure-tested those contracts by asking \"can "
        "a fresh implementer build the page from the spec alone?\" The V1 to V4 ladder "
        "is the record of how that gap closed (most visibly via the ABBackBar and "
        "ABPageHero split). Iteration 4 added the only thing the spec couldn't supply: "
        "state that crosses pages. Without it, Context-First is just words on a screen."
    )

    add_para(doc,
        "What I would do differently: I would start the App Target shell at Iteration "
        "2, not Iteration 4. Two of the three SPEC GAP categories (cross-page state, "
        "navigation, shared context) only surface when pages talk to each other. "
        "Waiting until Iteration 4 to discover them cost a full revision pass on "
        "Pages/. The lesson: spec-driven discipline is necessary but not sufficient. "
        "Cross-page state must be exercised early to close the loop."
    )

    add_label_value(doc, "Screenshot / Video Demo Link", "[VIDEO_LINK_TBD]: replace with YouTube unlisted link after recording.", value_italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
